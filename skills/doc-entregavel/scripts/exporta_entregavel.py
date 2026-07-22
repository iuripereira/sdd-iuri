#!/usr/bin/env python3
"""Exporta markdown -> PDF/DOCX assinável com capa parametrizada (skill doc-entregavel, ADR-0009).

Generalização do pipeline validado nos PRDs/contratos IMEX (imex-contratos):
- docx: pandoc (pypandoc) md->docx para o corpo; capa inserida via python-docx.
- pdf:  python-markdown -> HTML+CSS -> google-chrome --headless --print-to-pdf.

Imagens referenciadas no markdown (`![](docs/diagrams/arquitetura.svg)`) são embutidas
pelos dois pipelines — renderize os diagramas do doc-profile ANTES (mmdc/dbml-renderer;
PNG para docx, SVG ou PNG para pdf) e referencie-os no md de entrada.

Formatação padrão do entregável: corpo com linhas justificadas (títulos, tabelas e
código ficam à esquerda) e Sumário em página própria após a capa, no formato de
contrato (título pontilhado até o nº de página). No pdf os números vêm de duas
passadas de render (a 1ª mede em que página cada título cai — extração de texto por
pdftotext ou pypdf; sem nenhum dos dois, o Sumário sai sem números, com aviso). No
docx o Sumário é campo TOC nativo que o próprio Word preenche/atualiza ao abrir
(updateFields) — pontilhado e paginação são do Word.

Uso:
  exporta_entregavel.py {docx|pdf} <entrada.md> <saida> \
      --titulo "Documento de Requisitos do Produto (PRD)" \
      --projeto "Nome do Projeto" --versao 1.0 --data "20 de julho de 2026" \
      [--anexo "Anexo V ao Contrato ..."] [--local "Cidade/UF"] \
      [--assinatura "Razão Social — CONTRATANTE"] [--assinatura "... — CONTRATADO"]

Entregável é congelado: nova baseline -> novo arquivo versionado -> nova assinatura.
Nunca sobrescreva um entregável já enviado ao cliente.
"""
import argparse
import pathlib
import re
import subprocess
import sys
import tempfile

# Cambria = fonte dos contratos IMEX; Caladea = metricamente compatível; Noto Serif = fallback
CSS = '''
@page { size: 8.5in 11in; margin: 2.2cm 2.4cm; }
@page paisagem { size: 11in 8.5in; }
body { font-family: Cambria, Caladea, 'Noto Serif', Georgia, serif; font-size: 10.5pt; color: #000;
       text-align: justify; hyphens: auto; }
h1 { font-size: 13.5pt; color: #365F91; font-weight: bold; }
h2 { font-size: 12.5pt; color: #4F81BD; }
h3 { font-size: 11.5pt; color: #4F81BD; }
/* justificado é só para o corpo — título, tabela e código ficam à esquerda */
h1, h2, h3, h4, h5, h6, th, td, pre { text-align: left; }
/* Sumário no formato de contrato: título ..... nº de página */
.indice p { margin: 5pt 0; }
.indice a { color: #000; text-decoration: none; display: flex; align-items: baseline; }
.indice .dots { flex: 1; overflow: hidden; white-space: nowrap; margin: 0 3pt; }
.indice .dots::after { letter-spacing: 2pt; content: "................................................\
................................................................................................\
................................................................................................"; }
.ind-2 { margin-left: 1.2em; }
.ind-3 { margin-left: 2.4em; }
table { border-collapse: collapse; width: 100%; font-size: 9.5pt; break-inside: avoid; }
tr { break-inside: avoid; }
thead { display: table-header-group; }
th, td { border: 0.5pt solid #999; padding: 3pt 5pt; text-align: left; vertical-align: top; }
code, pre { font-family: 'Courier New', monospace; font-size: 9pt; }
pre { white-space: pre-wrap; }
img { max-width: 100%; break-inside: avoid; }
/* diagrama ocupa a própria página, centralizado; combine com .paisagem se for largo.
   O <p> que o md_in_html embrulha na imagem quebraria a cadeia de max-height — vira flex 100%. */
.fig-pagina { break-before: page; break-after: page; display: flex; align-items: center;
              justify-content: center; height: 9.2in; }
.fig-pagina > p { height: 100%; width: 100%; margin: 0; display: flex;
                  align-items: center; justify-content: center; }
.fig-pagina img { max-height: 100%; max-width: 100%; }
.paisagem { page: paisagem; }
.paisagem.fig-pagina, .paisagem .fig-pagina { height: 6.7in; }
blockquote { margin-left: 0.5cm; color: #333; }
.capa { text-align: center; margin-top: 3cm; }
.capa h1 { font-size: 26pt; color: #17365D; }
.assin { margin-top: 2.5cm; text-align: left; }
.linha { margin-top: 1.6cm; border-top: 0.75pt solid #000; width: 11cm; padding-top: 2pt; }
.quebra { page-break-after: always; }
'''


def le_markdown(caminho):
    md = pathlib.Path(caminho).read_text(encoding='utf-8')
    # TOC markdown (âncoras internas) vira ruído no documento final
    return re.sub(r'<!-- BEGIN TOC -->.*?<!-- END TOC -->', '', md, flags=re.S)


def _achata_toc(tokens):
    """toc_tokens (árvore) -> lista plana na ordem do documento."""
    plano = []
    for t in tokens:
        plano.append(t)
        plano.extend(_achata_toc(t['children']))
    return plano


def _texto_por_pagina(pdf):
    """Texto de cada página do pdf — pdftotext (poppler) ou pypdf; [] se nenhum existir."""
    from shutil import which
    if which('pdftotext'):
        with tempfile.TemporaryDirectory() as tmp:
            txt = pathlib.Path(tmp) / 'p.txt'
            subprocess.run(['pdftotext', str(pdf), str(txt)], check=True, capture_output=True)
            return txt.read_text(encoding='utf-8').split('\f')
    try:
        from pypdf import PdfReader
        return [p.extract_text() or '' for p in PdfReader(str(pdf)).pages]
    except ImportError:
        return []


def _paginas_dos_titulos(pdf, tokens):
    """{id do título: página física} — busca REVERSA com teto monotônico.

    Reversa porque todos os títulos também aparecem listados no próprio Sumário
    (páginas iniciais): procurando do fim para o início, a ocorrência encontrada
    é sempre a do título no corpo, nunca a da linha do Sumário.
    """
    paginas = _texto_por_pagina(pdf)
    if not paginas:
        print('aviso: sem pdftotext/pypdf — Sumário sai sem número de página '
              '(pip install pypdf)', file=sys.stderr)
        return {}
    def norm(s):
        return re.sub(r'\s+', '', s)
    textos = [norm(p) for p in paginas]
    pags, teto = {}, len(textos) - 1
    for t in reversed(tokens):
        for p in range(teto, -1, -1):
            if norm(t['name']) in textos[p]:
                pags[t['id']] = p + 1
                teto = p
                break
    return pags


def _html_sumario(tokens, pags):
    """Sumário no formato de contrato: título ..... nº (nº ausente = melhor esforço)."""
    import html as _html
    linhas = []
    for t in tokens:
        pg = pags.get(t['id'], '')
        linhas.append(f'<p class="ind-{t["level"]}"><a href="#{t["id"]}">'
                      f'<span>{_html.escape(t["name"])}</span>'
                      f'<span class="dots"></span><span>{pg}</span></a></p>')
    return ('<nav class="indice"><h1>Sumário</h1>' + '\n'.join(linhas)
            + '</nav><div class="quebra"></div>')


def exporta_pdf(args, md):
    import markdown
    # md_in_html: blocos <div class="paisagem"/"fig-pagina" markdown="1"> continuam processando markdown
    # toc: alimenta o Sumário com os títulos h1–h3 (âncoras clicáveis)
    mdx = markdown.Markdown(extensions=['tables', 'fenced_code', 'sane_lists', 'md_in_html', 'toc'],
                            extension_configs={'toc': {'toc_depth': '1-3'}})
    corpo = mdx.convert(md)
    tokens = _achata_toc(mdx.toc_tokens)
    assin = ''.join(f'<div class="linha"><b>{a}</b></div>' for a in args.assinatura)
    anexo = f'<p>{args.anexo}</p>' if args.anexo else ''
    local = (f'<p>{args.local}, ____ de ______________ de ____.</p>'
             if args.local else '')
    capa = f'''
<div class="capa">
  <h1>{args.titulo}</h1>
  <h2 style="color:#17365D">{args.projeto}</h2>
  <p><b>Versão {args.versao}, de {args.data}</b></p>
  {anexo}
  <div class="assin">{local}{assin}</div>
</div>
<div class="quebra"></div>
'''
    # <base> aponta para a pasta do md — imagens relativas resolvem de lá
    # lang=pt-BR habilita a hifenização do Chrome (hyphens: auto) no texto justificado
    base = pathlib.Path(args.entrada).resolve().parent.as_uri() + '/'

    def monta(sumario):
        return (f"<html lang='pt-BR'><head><meta charset='utf-8'><base href='{base}'>"
                f"<title>{args.projeto} v{args.versao}</title>"
                f"<style>{CSS}</style></head><body>{capa}{sumario}{corpo}</body></html>")

    with tempfile.TemporaryDirectory() as tmp:
        h = pathlib.Path(tmp) / 'doc.html'
        destino = pathlib.Path(args.saida)
        destino.parent.mkdir(parents=True, exist_ok=True)

        def imprime(alvo):
            subprocess.run(['google-chrome', '--headless=new', '--disable-gpu', '--no-sandbox',
                            '--no-pdf-header-footer', f'--print-to-pdf={alvo}',
                            str(h)], check=True, capture_output=True)

        # Passada 1: Sumário sem números (mesmo nº de linhas -> mesma paginação)
        # só para medir em que página física cada título cai.
        h.write_text(monta(_html_sumario(tokens, {})), encoding='utf-8')
        prova = pathlib.Path(tmp) / 'prova.pdf'
        imprime(prova.resolve())
        # Passada 2: Sumário definitivo com os números medidos.
        h.write_text(monta(_html_sumario(tokens, _paginas_dos_titulos(prova, tokens))),
                     encoding='utf-8')
        imprime(destino.resolve())


def _tabelas_sem_corte(d):
    """Linhas não quebram entre páginas; cabeçalho repete quando a tabela transborda.

    ponytail: docx não tem "tabela inteira numa página"; cantSplit por linha + tblHeader
    é o que o formato oferece — a quebra nunca corta uma linha ao meio.
    """
    from docx.oxml.ns import qn
    for t in d.tables:
        for i, row in enumerate(t.rows):
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(tr_pr.makeelement(qn('w:cantSplit'), {}))
            if i == 0:
                tr_pr.append(tr_pr.makeelement(qn('w:tblHeader'), {}))


def _justifica_corpo(d):
    """Parágrafos de corpo justificados; título/código/legenda ficam como estão.

    Só parágrafos de nível de documento (d.paragraphs exclui células de tabela)
    e sem alinhamento explícito — a capa, inserida depois, define o seu próprio.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    NAO_JUSTIFICA = ('Heading', 'Title', 'Subtitle', 'TOC', 'Caption', 'Source Code')
    for p in d.paragraphs:
        if p.alignment is None and not p.style.name.startswith(NAO_JUSTIFICA):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _campo_sumario(paragrafo):
    """Campo TOC nativo (h1–h3, hyperlinks) — o Word gera o Sumário ao abrir/F9,
    com o pontilhado e os números de página dele."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'Sumário — abra no Word e atualize os campos (F9) para preencher.'
    r.append(t)
    fld.append(r)
    paragrafo._p.append(fld)


def _atualiza_campos_ao_abrir(d):
    """settings.xml: updateFields=true — o Word oferece atualizar o TOC ao abrir."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    s = d.settings.element
    if s.find(qn('w:updateFields')) is None:
        u = OxmlElement('w:updateFields')
        u.set(qn('w:val'), 'true')
        s.append(u)


def exporta_docx(args, md):
    import pypandoc
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Pt

    C, E = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT
    with tempfile.TemporaryDirectory() as tmp:
        corpo = pathlib.Path(tmp) / 'corpo.docx'
        pypandoc.convert_text(
            md, 'docx', format='gfm', outputfile=str(corpo),
            extra_args=['--resource-path', str(pathlib.Path(args.entrada).resolve().parent)])
        d = Document(str(corpo))
        _justifica_corpo(d)

        linhas = [(args.titulo, 22, True, C),
                  (args.projeto, 16, True, C),
                  (f'Versão {args.versao}, de {args.data}', 12, True, C)]
        if args.anexo:
            linhas.append((args.anexo, 11, False, C))
        if args.local:
            linhas += [('', 11, False, E),
                       (f'{args.local}, ____ de ______________ de ____.', 11, False, E)]
        for a in args.assinatura:
            linhas += [('', 11, False, E),
                       ('_________________________________________________', 11, False, E),
                       (a, 11, True, E)]
        # insert_paragraph_before(primeiro) insere sempre imediatamente antes de `primeiro`,
        # ou seja, APÓS as linhas já inseridas — a ordem de iteração é a ordem final.
        primeiro = d.paragraphs[0]
        for texto, tam, negrito, alinh in linhas:
            p = primeiro.insert_paragraph_before('')
            r = p.add_run(texto)
            r.bold = negrito
            r.font.size = Pt(tam)
            p.alignment = alinh
        quebra = primeiro.insert_paragraph_before('')
        quebra.add_run().add_break(WD_BREAK.PAGE)

        # Sumário em página própria, entre a capa e o corpo
        tit = primeiro.insert_paragraph_before('')
        r = tit.add_run('Sumário')
        r.bold = True
        r.font.size = Pt(14)
        _campo_sumario(primeiro.insert_paragraph_before(''))
        quebra2 = primeiro.insert_paragraph_before('')
        quebra2.add_run().add_break(WD_BREAK.PAGE)
        _atualiza_campos_ao_abrir(d)

        _tabelas_sem_corte(d)
        destino = pathlib.Path(args.saida)
        destino.parent.mkdir(parents=True, exist_ok=True)
        d.save(str(destino))


def selftest():
    """Fixture mínima md -> docx (e pdf, se houver chrome); acusa saída vazia/ausente."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp = pathlib.Path(tmp)
        (tmp / 'doc.md').write_text('# Título\n\ncorpo\n\n## Seção\n\n| a | b |\n|---|---|\n| 1 | 2 |\n',
                                    encoding='utf-8')
        base = ['--titulo', 'T', '--projeto', 'P', '--versao', '0.0', '--data', 'hoje',
                '--local', 'X/Y', '--assinatura', 'A — CONTRATANTE']
        for fmt in ('docx', 'pdf'):
            if fmt == 'pdf':
                from shutil import which
                if not which('google-chrome'):
                    print('selftest: pdf PULADO (google-chrome ausente)')
                    continue
            saida = tmp / f'doc.{fmt}'
            main([fmt, str(tmp / 'doc.md'), str(saida)] + base)
            assert saida.exists() and saida.stat().st_size > 0, f'{fmt}: saída vazia'
            if fmt == 'docx':
                from docx import Document
                d = Document(str(saida))
                xml = d.element.xml
                assert 'TOC' in xml and 'Sumário' in xml, 'docx: campo de sumário ausente'
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                assert any(p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
                           for p in d.paragraphs), 'docx: corpo não justificado'
            if fmt == 'pdf':
                paginas = _texto_por_pagina(saida)
                if paginas:  # extração disponível: Sumário na pág. 2 com linha pontilhada e nº
                    assert 'Sumário' in paginas[1], 'pdf: página de Sumário ausente'
                    # pdftotext extrai o pontilhado com espaços entre os pontos (". . . . 3")
                    assert re.search(r'(\.\s*){4,}\d', paginas[1]), 'pdf: sumário sem nº de página'
                else:
                    print('selftest: nº de página do sumário NÃO VERIFICADO (sem pdftotext/pypdf)')
            print(f'selftest: {fmt} OK ({saida.stat().st_size} bytes)')
    print('selftest: OK')


def main(argv):
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument('formato', choices=['docx', 'pdf'])
    ap.add_argument('entrada')
    ap.add_argument('saida')
    ap.add_argument('--titulo', required=True)
    ap.add_argument('--projeto', required=True)
    ap.add_argument('--versao', required=True)
    ap.add_argument('--data', required=True)
    ap.add_argument('--anexo', default='')
    ap.add_argument('--local', default='')
    ap.add_argument('--assinatura', action='append', default=[])
    args = ap.parse_args(argv)
    md = le_markdown(args.entrada)
    (exporta_docx if args.formato == 'docx' else exporta_pdf)(args, md)
    print('OK', args.saida)


if __name__ == '__main__':
    if '--selftest' in sys.argv:
        selftest()
    else:
        main(sys.argv[1:])
